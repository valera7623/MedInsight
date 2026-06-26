"""Generate clinical DOCX reports with python-docx."""

from __future__ import annotations

import logging
import re
from datetime import date, datetime
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from sqlalchemy.orm import Session

from app.config import settings
from app.models import Department, DicomStudy, Document, Patient, Prediction, Tenant
from app.services.docx_templates import (
    DEFAULT_PATIENT_CARD_SECTIONS,
    GENDER_LABELS,
    TEMPLATE_CLINICAL_SUMMARY,
    TEMPLATE_DICOM_REPORT,
    TEMPLATE_LAB_REPORT,
    TEMPLATE_PATIENT_CARD,
)
from app.services.export_report import collect_patient_export_clinical_data
from app.services.extractor import diagnoses_from_parsed_data, labs_dict_to_list, medications_from_parsed_data
from app.templates.docx.patient_card_styles import (
    COLOR_MUTED,
    COLOR_PRIMARY,
    COLOR_WATERMARK,
    FONT_BODY_SIZE,
    FONT_HEADING_SIZE,
    FONT_NAME,
    FONT_SUBHEADING_SIZE,
    FONT_TABLE_SIZE,
    FONT_TITLE_SIZE,
    LINE_SPACING,
    PAGE_MARGINS,
    TABLE_STYLE,
)

logger = logging.getLogger(__name__)

_ICD_PREFIX = re.compile(r"^([A-ZА-Я]\d{2}(?:\.\d{1,2})?)\s*(?:\((.+)\))?$", re.IGNORECASE)


class DocxGenerator:
    """Build Word documents for patient cards and clinical summaries."""

    def __init__(self, db: Session | None = None) -> None:
        self.db = db

    # ------------------------------------------------------------------ helpers

    @staticmethod
    def _apply_page_margins(doc: Document) -> None:
        for section in doc.sections:
            section.left_margin = PAGE_MARGINS["left"]
            section.right_margin = PAGE_MARGINS["right"]
            section.top_margin = PAGE_MARGINS["top"]
            section.bottom_margin = PAGE_MARGINS["bottom"]

    @staticmethod
    def _style_run(run, *, size: Pt = FONT_BODY_SIZE, bold: bool = False, color=COLOR_PRIMARY) -> None:
        run.font.name = FONT_NAME
        run.font.size = size
        run.font.bold = bold
        if color is not None:
            run.font.color.rgb = color

    @classmethod
    def add_paragraph_with_style(
        cls,
        doc: Document,
        text: str,
        style: str = "Normal",
        *,
        bold: bool = False,
        align: WD_ALIGN_PARAGRAPH | None = None,
    ):
        paragraph = doc.add_paragraph(style=style)
        run = paragraph.add_run(text)
        cls._style_run(run, bold=bold, color=COLOR_PRIMARY if bold else None)
        paragraph.paragraph_format.line_spacing = LINE_SPACING
        if align is not None:
            paragraph.alignment = align
        return paragraph

    @classmethod
    def add_heading(cls, doc: Document, text: str, level: int = 1):
        heading = doc.add_heading(text, level=level)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        size = FONT_HEADING_SIZE if level == 1 else FONT_SUBHEADING_SIZE
        for run in heading.runs:
            cls._style_run(run, size=size, bold=True)
        return heading

    @classmethod
    def add_table(cls, doc: Document, data: list[list[Any]], headers: list[str], title: str | None = None):
        if title:
            cls.add_paragraph_with_style(doc, title, bold=True)

        table = doc.add_table(rows=1, cols=len(headers))
        table.style = TABLE_STYLE
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        header_cells = table.rows[0].cells
        for index, header in enumerate(headers):
            header_cells[index].text = str(header)
            if header_cells[index].paragraphs[0].runs:
                cls._style_run(
                    header_cells[index].paragraphs[0].runs[0],
                    size=FONT_TABLE_SIZE,
                    bold=True,
                    color=None,
                )
            header_cells[index].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for row_data in data:
            row_cells = table.add_row().cells
            for index, cell_data in enumerate(row_data):
                row_cells[index].text = str(cell_data)
                if row_cells[index].paragraphs[0].runs:
                    cls._style_run(row_cells[index].paragraphs[0].runs[0], size=FONT_TABLE_SIZE, color=None)

        if table.columns:
            for cell in table.columns[0].cells:
                cell.width = Cm(4)
        return table

    @classmethod
    def add_header_footer(cls, doc: Document, header_text: str, footer_text: str) -> None:
        for section in doc.sections:
            header = section.header
            header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            header_para.text = header_text
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if header_para.runs:
                cls._style_run(header_para.runs[0], size=Pt(9), bold=True, color=COLOR_MUTED)

            footer = section.footer
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.text = footer_text
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if footer_para.runs:
                cls._style_run(footer_para.runs[0], size=Pt(9), color=COLOR_MUTED)

    @classmethod
    def add_watermark(cls, doc: Document, text: str) -> None:
        """Light watermark in the document header."""
        for section in doc.sections:
            header = section.header
            paragraph = header.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(text)
            run.font.name = FONT_NAME
            run.font.size = Pt(36)
            run.font.color.rgb = COLOR_WATERMARK

    @classmethod
    def add_lab_results_table(cls, doc: Document, lab_results: list[dict[str, Any]]):
        headers = TEMPLATE_LAB_REPORT["table_headers"]
        rows: list[list[str]] = []
        for result in lab_results:
            status = result.get("status") or _lab_status_label(result)
            rows.append(
                [
                    result.get("name", ""),
                    result.get("value", ""),
                    result.get("reference", ""),
                    status,
                ]
            )
        return cls.add_table(doc, rows, headers)

    @classmethod
    def add_dicom_info(cls, doc: Document, dicom_study: dict[str, Any]) -> None:
        cls.add_paragraph_with_style(doc, "DICOM-исследование", bold=True)
        for label, key in (
            ("Модальность", "modality"),
            ("Область", "body_part"),
            ("Описание", "study_description"),
            ("Дата исследования", "study_date"),
            ("Количество серий", "num_series"),
            ("Количество кадров", "num_instances"),
        ):
            value = dicom_study.get(key, "—")
            cls.add_paragraph_with_style(doc, f"{label}: {value}")
        if dicom_study.get("radiology_impression"):
            cls.add_paragraph_with_style(doc, f"Заключение: {dicom_study['radiology_impression']}")
        doc.add_paragraph("")

    # ------------------------------------------------------------------ builders

    @classmethod
    def build_patient_card_bytes(
        cls,
        *,
        patient_data: dict[str, Any],
        lab_results: list[dict[str, Any]],
        diagnoses: list[dict[str, Any]],
        medications: list[dict[str, Any]],
        predictions: list[dict[str, Any]],
        dicom_studies: list[dict[str, Any]],
        anamnesis: list[str] | None = None,
        operations: list[str] | None = None,
        imaging: list[str] | None = None,
        recommendations: list[str] | None = None,
        sections: list[str] | None = None,
        header_text: str = "",
        footer_text: str = "",
        watermark: str | None = None,
    ) -> BytesIO:
        enabled = set(sections or DEFAULT_PATIENT_CARD_SECTIONS)
        doc = Document()
        cls._apply_page_margins(doc)

        if header_text or footer_text:
            cls.add_header_footer(doc, header_text, footer_text)
        if watermark:
            cls.add_watermark(doc, watermark)

        title = doc.add_heading(TEMPLATE_PATIENT_CARD["title"], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            cls._style_run(run, size=FONT_TITLE_SIZE, bold=True)

        cls.add_paragraph_with_style(
            doc,
            f"Дата формирования: {datetime.now().strftime('%d.%m.%Y %H:%M')}",
        )
        doc.add_paragraph("")

        section_no = 1

        if "patient" in enabled:
            cls.add_heading(doc, f"{section_no}. ИНФОРМАЦИЯ О ПАЦИЕНТЕ", 1)
            for label, key in (
                ("ФИО", "full_name"),
                ("Дата рождения", "birth_date"),
                ("Пол", "gender"),
                ("Телефон", "phone"),
                ("Email", "email"),
                ("Отделение", "department"),
            ):
                cls.add_paragraph_with_style(doc, f"{label}: {patient_data.get(key, '—')}")
            doc.add_paragraph("")
            section_no += 1

        if "anamnesis" in enabled and anamnesis:
            cls.add_heading(doc, f"{section_no}. АНАМНЕЗ", 1)
            for item in anamnesis:
                doc.add_paragraph(f"• {item}", style="List Bullet")
            doc.add_paragraph("")
            section_no += 1

        if "diagnoses" in enabled and diagnoses:
            cls.add_heading(doc, f"{section_no}. ДИАГНОЗЫ", 1)
            for diag in diagnoses:
                code = diag.get("code", "")
                name = diag.get("name", "")
                line = f"{code} — {name}".strip(" —")
                doc.add_paragraph(f"• {line}", style="List Bullet")
            doc.add_paragraph("")
            section_no += 1

        if "lab" in enabled and lab_results:
            cls.add_heading(doc, f"{section_no}. ЛАБОРАТОРНЫЕ АНАЛИЗЫ", 1)
            cls.add_lab_results_table(doc, lab_results)
            doc.add_paragraph("")
            section_no += 1

        if "medications" in enabled and medications:
            cls.add_heading(doc, f"{section_no}. ЛЕКАРСТВА", 1)
            headers = ["Название", "Дозировка", "Частота", "Назначен"]
            rows = [
                [
                    med.get("name", ""),
                    med.get("dosage", "—"),
                    med.get("frequency", "—"),
                    med.get("prescribed_date", "—"),
                ]
                for med in medications
            ]
            cls.add_table(doc, rows, headers)
            doc.add_paragraph("")
            section_no += 1

        if "predictions" in enabled and predictions:
            cls.add_heading(doc, f"{section_no}. ПРОГНОЗЫ", 1)
            for pred in predictions:
                cls.add_paragraph_with_style(
                    doc,
                    f"• {pred.get('type', 'Прогноз')}: риск {pred.get('risk', 0)}%",
                    bold=False,
                )
                factors = pred.get("factors") or []
                if factors:
                    cls.add_paragraph_with_style(doc, f"  Факторы: {', '.join(factors)}")
                recs = pred.get("recommendations") or []
                if recs:
                    cls.add_paragraph_with_style(doc, f"  Рекомендации: {', '.join(recs)}")
            doc.add_paragraph("")
            section_no += 1

        if "dicom" in enabled and dicom_studies:
            cls.add_heading(doc, f"{section_no}. DICOM-ИССЛЕДОВАНИЯ", 1)
            for study in dicom_studies:
                cls.add_dicom_info(doc, study)
            section_no += 1

        if "conclusion" in enabled:
            cls.add_heading(doc, f"{section_no}. ЗАКЛЮЧЕНИЕ И РЕКОМЕНДАЦИИ", 1)
            if recommendations:
                for rec in recommendations:
                    doc.add_paragraph(f"• {rec}", style="List Bullet")
            else:
                cls.add_paragraph_with_style(doc, "_________________________________________________________________")
                doc.add_paragraph("")
                cls.add_paragraph_with_style(doc, "_________________________________________________________________")
            doc.add_paragraph("")

        if operations and "operations" in enabled:
            pass  # reserved for clinical summary template

        if imaging and "imaging" in enabled:
            pass

        doc.add_paragraph("")
        cls.add_paragraph_with_style(doc, "Лечащий врач: ____________________")
        cls.add_paragraph_with_style(doc, f"Дата: {datetime.now().strftime('%d.%m.%Y')}")

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    def generate_patient_card(self, patient_id: int, options: dict | None = None) -> BytesIO:
        if self.db is None:
            raise ValueError("Database session is required for generate_patient_card")
        context = collect_patient_docx_context(self.db, patient_id, options or {})
        return self.build_patient_card_bytes(**context)

    def generate_lab_report(self, patient_id: int, lab_results: list[dict[str, Any]]) -> BytesIO:
        if self.db is None:
            raise ValueError("Database session is required")
        patient = self.db.query(Patient).filter(Patient.id == patient_id).first()
        if not patient:
            raise ValueError(f"Patient {patient_id} not found")

        patient_data = _patient_to_dict(patient, self.db)
        return self.build_patient_card_bytes(
            patient_data=patient_data,
            lab_results=lab_results,
            diagnoses=[],
            medications=[],
            predictions=[],
            dicom_studies=[],
            sections=["patient", "lab"],
            header_text=TEMPLATE_LAB_REPORT["title"],
            footer_text=settings.DOCX_WATERMARK,
            watermark=settings.DOCX_WATERMARK,
        )

    def generate_clinical_summary(self, patient_id: int, options: dict | None = None) -> BytesIO:
        if self.db is None:
            raise ValueError("Database session is required")
        context = collect_patient_docx_context(self.db, patient_id, options or {})
        return self.build_patient_card_bytes(
            **context,
            sections=["patient", "anamnesis", "diagnoses", "conclusion"],
            header_text=TEMPLATE_CLINICAL_SUMMARY["title"],
            footer_text=settings.DOCX_WATERMARK,
            watermark=settings.DOCX_WATERMARK,
        )

    def generate_dicom_report(self, dicom_study_uid: str) -> BytesIO:
        if self.db is None:
            raise ValueError("Database session is required")
        study = self.db.query(DicomStudy).filter(DicomStudy.study_uid == dicom_study_uid).first()
        if not study:
            raise ValueError(f"DICOM study {dicom_study_uid} not found")

        patient = self.db.query(Patient).filter(Patient.id == study.patient_id).first()
        patient_data = _patient_to_dict(patient, self.db) if patient else {}

        doc = Document()
        self._apply_page_margins(doc)
        self.add_header_footer(doc, TEMPLATE_DICOM_REPORT["title"], settings.DOCX_WATERMARK)
        self.add_watermark(doc, settings.DOCX_WATERMARK)

        title = doc.add_heading(TEMPLATE_DICOM_REPORT["title"], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            self._style_run(run, size=FONT_TITLE_SIZE, bold=True)

        self.add_heading(doc, "1. ПАЦИЕНТ", 1)
        self.add_paragraph_with_style(doc, f"ФИО: {patient_data.get('full_name', '—')}")
        self.add_paragraph_with_style(doc, f"ID пациента: {study.patient_id}")
        doc.add_paragraph("")

        self.add_heading(doc, "2. ПАРАМЕТРЫ ИССЛЕДОВАНИЯ", 1)
        self.add_dicom_info(doc, _dicom_study_to_dict(study))

        if study.radiology_findings:
            self.add_heading(doc, "3. ОПИСАНИЕ", 1)
            for item in study.radiology_findings:
                if isinstance(item, str):
                    self.add_paragraph_with_style(doc, f"• {item}")
                elif isinstance(item, dict):
                    self.add_paragraph_with_style(doc, f"• {item.get('text', item)}")

        if study.radiology_impression:
            self.add_heading(doc, "4. ЗАКЛЮЧЕНИЕ", 1)
            self.add_paragraph_with_style(doc, study.radiology_impression)

        if study.extracted_measurements:
            self.add_heading(doc, "5. ИЗМЕРЕНИЯ", 1)
            measurements = study.extracted_measurements
            if isinstance(measurements, dict):
                rows = [[key, str(value)] for key, value in measurements.items()]
                self.add_table(doc, rows, ["Параметр", "Значение"])

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer


# ---------------------------------------------------------------------------
# Data collection
# ---------------------------------------------------------------------------


def _lab_status_label(entry: dict[str, Any]) -> str:
    if entry.get("status"):
        status = str(entry["status"]).casefold()
        if status in {"повышен", "high", "h"}:
            return "↑"
        if status in {"понижен", "low", "l"}:
            return "↓"
        return entry["status"]
    if entry.get("abnormal"):
        return "отклонение"
    return "норма"


def _parse_diagnosis(text: str) -> dict[str, str]:
    match = _ICD_PREFIX.match(text.strip())
    if match:
        return {"code": match.group(1).upper(), "name": (match.group(2) or match.group(1)).strip()}
    return {"code": "", "name": text.strip()}


def _format_date(value: date | datetime | str | None) -> str:
    if value is None:
        return "—"
    if isinstance(value, datetime):
        return value.strftime("%d.%m.%Y")
    if isinstance(value, date):
        return value.strftime("%d.%m.%Y")
    return str(value)


def _patient_to_dict(patient: Patient, db: Session) -> dict[str, Any]:
    department_name = "—"
    if patient.department_id:
        dept = db.query(Department).filter(Department.id == patient.department_id).first()
        if dept:
            department_name = dept.name

    parts = [patient.last_name, patient.first_name]
    if patient.middle_name:
        parts.append(patient.middle_name)

    return {
        "full_name": " ".join(parts),
        "birth_date": _format_date(patient.birth_date),
        "gender": GENDER_LABELS.get(patient.gender, patient.gender),
        "phone": patient.phone or "—",
        "email": patient.email or "—",
        "department": department_name,
    }


def _dicom_study_to_dict(study: DicomStudy) -> dict[str, Any]:
    return {
        "modality": study.modality or "—",
        "body_part": study.body_part or "—",
        "study_description": study.study_description or "—",
        "study_date": _format_date(study.study_date),
        "num_series": study.num_series,
        "num_instances": study.num_instances,
        "radiology_impression": study.radiology_impression,
    }


def _merge_clinical_from_documents(documents: list[Document]) -> dict[str, Any]:
    diagnoses: set[str] = set()
    medications: set[str] = set()
    anamnesis: set[str] = set()
    operations: set[str] = set()
    imaging: set[str] = set()
    labs: dict[str, dict[str, Any]] = {}

    for doc in documents:
        parsed = doc.parsed_data or {}
        if doc.status != "parsed" or not parsed:
            continue
        diagnoses.update(diagnoses_from_parsed_data(parsed))
        medications.update(medications_from_parsed_data(parsed))
        for item in parsed.get("anamnesis") or []:
            if item:
                anamnesis.add(str(item))
        for item in parsed.get("operations") or []:
            if item:
                operations.add(str(item))
        for item in parsed.get("imaging_conclusions") or []:
            if item:
                imaging.add(str(item))
        for lab in labs_dict_to_list(parsed.get("lab_results")):
            key = (lab.get("name") or "").casefold()
            if key:
                labs[key] = lab

    return {
        "diagnoses": [_parse_diagnosis(d) for d in sorted(diagnoses)],
        "medications": [{"name": m, "dosage": "—", "frequency": "—", "prescribed_date": "—"} for m in sorted(medications, key=str.casefold)],
        "anamnesis": sorted(anamnesis, key=str.casefold),
        "operations": sorted(operations, key=str.casefold),
        "imaging": sorted(imaging, key=str.casefold),
        "lab_results": [
            {**lab, "status": _lab_status_label(lab)}
            for lab in sorted(labs.values(), key=lambda x: (x.get("name") or "").casefold())
        ],
    }


def _predictions_from_db(predictions: list[Prediction]) -> tuple[list[dict[str, Any]], list[str]]:
    if not predictions:
        return [], []

    latest = sorted(predictions, key=lambda p: p.created_at, reverse=True)[0]
    payload = latest.prediction or {}
    readmission = round(float(payload.get("readmission_risk") or 0))
    complication = round(float(payload.get("complication_risk") or 0))
    factors = [str(f) for f in (payload.get("factors") or [])]
    recommendations = [str(r) for r in (payload.get("recommendations") or [])]

    items = [
        {"type": "Риск реадмиссии", "risk": readmission, "factors": factors, "recommendations": recommendations},
        {"type": "Риск осложнений", "risk": complication, "factors": factors, "recommendations": recommendations},
    ]
    return items, recommendations


def collect_patient_docx_context(db: Session, patient_id: int, options: dict[str, Any]) -> dict[str, Any]:
    patient = db.query(Patient).filter(Patient.id == patient_id).first()
    if not patient:
        raise ValueError(f"Patient {patient_id} not found")

    tenant = db.query(Tenant).filter(Tenant.id == patient.tenant_id).first()
    clinic_name = tenant.name if tenant else "MedInsight"

    documents = (
        db.query(Document)
        .filter(Document.patient_id == patient_id, Document.tenant_id == patient.tenant_id)
        .order_by(Document.created_at.desc())
        .all()
    )
    clinical = _merge_clinical_from_documents(documents)

    # Fallback diagnoses/medications from export helper (handles legacy parsed_data).
    export_diagnoses, export_meds, discharge_blocks = collect_patient_export_clinical_data(documents)
    if export_diagnoses and not clinical["diagnoses"]:
        clinical["diagnoses"] = [_parse_diagnosis(d) for d in export_diagnoses]
    if export_meds and not clinical["medications"]:
        clinical["medications"] = [
            {"name": m, "dosage": "—", "frequency": "—", "prescribed_date": "—"}
            for m in export_meds
        ]
    if discharge_blocks and not clinical["anamnesis"]:
        clinical["anamnesis"] = [text[:500] for _, text in discharge_blocks[:3]]

    predictions_db = (
        db.query(Prediction)
        .filter(Prediction.patient_id == patient_id)
        .order_by(Prediction.created_at.desc())
        .all()
    )
    prediction_items, recommendations = _predictions_from_db(predictions_db)

    dicom_studies: list[dict[str, Any]] = []
    if settings.DICOM_ENABLED:
        studies = (
            db.query(DicomStudy)
            .filter(DicomStudy.patient_id == patient_id, DicomStudy.status == "ready")
            .order_by(DicomStudy.study_date.desc(), DicomStudy.created_at.desc())
            .all()
        )
        dicom_studies = [_dicom_study_to_dict(s) for s in studies]

    sections = options.get("sections") or DEFAULT_PATIENT_CARD_SECTIONS
    watermark = options.get("watermark") or settings.DOCX_WATERMARK

    return {
        "patient_data": _patient_to_dict(patient, db),
        "lab_results": clinical["lab_results"],
        "diagnoses": clinical["diagnoses"],
        "medications": clinical["medications"],
        "predictions": prediction_items,
        "dicom_studies": dicom_studies,
        "anamnesis": clinical["anamnesis"],
        "operations": clinical["operations"],
        "imaging": clinical["imaging"],
        "recommendations": recommendations,
        "sections": sections,
        "header_text": clinic_name,
        "footer_text": f"{clinic_name} · {datetime.now().strftime('%d.%m.%Y')}",
        "watermark": watermark,
    }


def save_docx_to_patient_reports(patient_id: int, buffer: BytesIO, *, suffix: str = "patient_card") -> str:
    """Persist generated DOCX under storage/reports/{patient_id}/."""
    base = Path(settings.DOCX_REPORTS_DIR) / str(patient_id)
    base.mkdir(parents=True, exist_ok=True)
    filename = f"{suffix}_{datetime.utcnow():%Y%m%d_%H%M%S}.docx"
    path = base / filename
    path.write_bytes(buffer.getvalue())
    logger.info("DOCX saved: %s", path)
    return str(path)
