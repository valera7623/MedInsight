#!/usr/bin/env python3
"""End-to-end test: patient → document upload → async parse → GPT prediction."""

import json
import sys
import time
from pathlib import Path

import httpx

BASE_URL = sys.argv[1] if len(sys.argv) > 1 else "http://localhost:8000"
EMAIL = "test_predict@clinic.ru"
PASSWORD = "test123456"


def main() -> None:
    client = httpx.Client(base_url=BASE_URL, timeout=120.0)

    print(f"=== MedInsight Phase 2 Test ({BASE_URL}) ===\n")

    # Register (ignore if exists)
    client.post(
        "/api/auth/register",
        json={
            "email": EMAIL,
            "password": PASSWORD,
            "full_name": "Test Predict Doctor",
            "role": "doctor",
        },
    )

    # Login
    login_res = client.post(
        "/api/auth/login",
        json={"email": EMAIL, "password": PASSWORD},
    )
    login_res.raise_for_status()
    token = login_res.json()["access_token"]
    headers = {"Authorization": f"Bearer {token}"}
    print("✓ Authenticated")

    # Create patient
    patient_res = client.post(
        "/api/patients",
        headers=headers,
        json={
            "first_name": "Тест",
            "last_name": "Прогнозов",
            "birth_date": "1970-05-20",
            "gender": "M",
            "phone": "+79001112233",
        },
    )
    patient_res.raise_for_status()
    patient_id = patient_res.json()["id"]
    print(f"✓ Patient created: id={patient_id}")

    # Create minimal test document (plain text saved as .docx-like content via upload)
    # Use a simple PDF or create a minimal docx - for test we'll create a tiny docx
    try:
        from docx import Document as DocxDocument

        test_file = Path("/tmp/medinsight_test.docx")
        doc = DocxDocument()
        doc.add_paragraph("Диагноз: J45.0 Бронхиальная астма. I10 Гипертензия.")
        doc.add_paragraph("Назначено: Амоксициллин, Парацетамол, Эналаприл.")
        doc.save(str(test_file))
    except ImportError:
        print("python-docx not installed — skipping document upload test")
        test_file = None

    doc_id = None
    if test_file and test_file.exists():
        with open(test_file, "rb") as f:
            upload_res = client.post(
                "/api/documents/upload",
                headers=headers,
                data={"patient_id": str(patient_id), "document_type": "discharge"},
                files={"file": ("test_discharge.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            )
        upload_res.raise_for_status()
        doc_id = upload_res.json()["id"]
        print(f"✓ Document uploaded: id={doc_id}, status={upload_res.json()['status']}")

        # Poll document parsing
        print("  Waiting for async parse...")
        for _ in range(30):
            doc_res = client.get(f"/api/documents/{doc_id}", headers=headers)
            doc_res.raise_for_status()
            doc_data = doc_res.json()
            if doc_data["status"] == "parsed":
                print(f"✓ Document parsed: {len(doc_data.get('parsed_data', {}).get('diagnoses', []))} diagnoses")
                break
            if doc_data["status"] == "failed":
                print(f"✗ Parse failed: {doc_data.get('parsed_data')}")
                break
            time.sleep(2)
        else:
            print("✗ Parse timeout")

    # Start prediction
    predict_res = client.post(f"/api/analytics/predict/{patient_id}", headers=headers)
    predict_res.raise_for_status()
    job_id = predict_res.json()["job_id"]
    print(f"✓ Prediction job started: job_id={job_id}")

    # Poll prediction status
    print("  Waiting for async prediction...")
    result = None
    for _ in range(30):
        status_res = client.get(f"/api/analytics/predict/status/{job_id}", headers=headers)
        status_res.raise_for_status()
        status_data = status_res.json()
        if status_data["status"] == "completed":
            result = status_data["result"]
            break
        if status_data["status"] == "failed":
            print(f"✗ Prediction failed: {status_data.get('error')}")
            sys.exit(1)
        time.sleep(2)
    else:
        print("✗ Prediction timeout")
        sys.exit(1)

    print("\n=== Prediction Result ===")
    print(json.dumps(result, ensure_ascii=False, indent=2))

    # List predictions
    preds_res = client.get(f"/api/analytics/predictions/{patient_id}", headers=headers)
    preds_res.raise_for_status()
    print(f"\n✓ Total predictions: {len(preds_res.json()['predictions'])}")

    # Insights
    insights_res = client.post(f"/api/analytics/insights/{patient_id}", headers=headers)
    insights_res.raise_for_status()
    insights = insights_res.json()
    print("\n=== AI Insights ===")
    print(insights.get("insights", ""))
    print("Recommendations:", insights.get("recommendations", []))

    print("\n=== Test completed successfully ===")


if __name__ == "__main__":
    main()
